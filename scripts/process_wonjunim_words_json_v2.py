import os
import glob
import json
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

# API Key from environment variable
API_KEY = os.environ.get("OPENAI_API_KEY")

# Paths
INPUT_DIR = r"d:\페르미아\가람컨텐츠팀장\원주님말씀.json\2024-2025"
OUTPUT_DIR = r"d:\페르미아\가람컨텐츠팀장\원주님말씀\2024-2025.new"

# Model Configuration
MODEL_NAME = "gpt-4o"

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Initialize OpenAI client
client = OpenAI(api_key=API_KEY)

def create_docx(filename, content_text):
    """
    Creates a DOCX file from the processed text.
    Assumes the text is formatted with Markdown-style headers (#, ##).
    """
    doc = Document()
    
    # Set default style (optional customization)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    lines = content_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Title or H1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # H3
            doc.add_heading(line[4:], level=3)
        else:
            # Paragraph
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.6

    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    print(f"Saved: {output_path}")

def process_file(json_path):
    base_name = os.path.basename(json_path)
    # Remove .json extension and replace with .docx
    docx_name = base_name.replace('.json', '.docx')
    output_path = os.path.join(OUTPUT_DIR, docx_name)

    if os.path.exists(output_path):
        print(f"Skipping (already exists): {docx_name}")
        return

    print(f"Processing: {base_name}")

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        # Extract text segments
        segments = data.get('segments', [])
        full_text = ""
        for seg in segments:
            text = seg.get('text', '')
            full_text += text + " "
        
        if not full_text.strip():
            print(f"Warning: No text found in {base_name}")
            return

        # Prepare prompt
        system_prompt = (
            "당신은 전문 편집자입니다. 주어진 녹취록은 스승 '원주님'과 제자들의 대화입니다. "
            "녹취록에서 **원주님의 가르침과 말씀** 부분을 중심으로 내용을 추출하여, "
            "읽기 쉽고 품위 있는 **문어체**로 변환하십시오. "
            "대화 형식이 아닌, 완성된 글(에세이 또는 강의록) 형태로 다듬어야 합니다. "
            "내용의 흐름에 따라 적절한 **소제목**(블로그나 책의 소제목 스타일)을 붙여 섹션을 구분해 주십시오. "
            "출력 형식은 Markdown 헤더(#, ##)를 사용하여 구조화해 주십시오. "
            "제자의 질문은 문맥상 꼭 필요한 경우에만 서술에 자연스럽게 녹여내고, "
            "불필요한 추임새나 반복은 제거하십시오."
        )

        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"다음 녹취록을 문어체로 변환하고 소제목을 붙여 정리해 주세요:\n\n{full_text}"}
            ],
            temperature=0.3,
            max_tokens=4000 
        )

        result_text = response.choices[0].message.content
        create_docx(docx_name, result_text)

    except Exception as e:
        print(f"Error processing {base_name}: {e}")

def main():
    json_files = sorted(glob.glob(os.path.join(INPUT_DIR, "*.json")))
    print(f"Found {len(json_files)} JSON files.")

    for json_file in json_files:
        process_file(json_file)
        # Sleep briefly to avoid rate limits if necessary, though gpt-4o limits are high
        time.sleep(1) 

if __name__ == "__main__":
    main()
