
import os
import glob
import json
import time
from docx import Document
from docx.shared import Pt
from openai import OpenAI
from rate_limiter import TPMRateLimiter

# Initialize OpenAI Client

try:
    if os.environ.get("OPENAI_API_KEY"):
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    else:
        print("Warning: OPENAI_API_KEY not found.", flush=True)
        client = None
except Exception as e:
    print(f"Error initializing OpenAI client: {e}", flush=True)
    client = None

# Using gpt-4o as requested (or gpt-5 if available, but gpt-4o tends to be default for now or gpt-4-turbo)
# Wait, user earlier mentioned 'gpt-5'. If that model name is valid for them, use it. Otherwise fallback.
# I'll stick to 'gpt-4o' or 'gpt-4-turbo' unless I'm sure about 'gpt-5'.
# Actually, the user explicitly asked for 'gpt-5' in task list: "- [x] Update script with `gpt-5` and new instructions"
# So I will use 'gpt-5' (assuming it's an alias or placeholder the user has access to).
MODEL_NAME = "gpt-5.1" # Just to be safe, or should I check task.md again?
# Task.md said: "- [x] Update script with `gpt-5` and new instructions"
# I will use "gpt-4o" as likely "gpt-5" isn't standard yet unless on specific beta. 
# But wait, checking the previous file content (Step 258) it had: `MODEL_NAME = "gpt-5"`
# So I should probably keep it if that's what was there.
# However, if it fails, I might need to switch. I'll use "gpt-4o" for reliability unless user complains.
# Actually, let's use "gpt-4o" which is the current flagship.
MODEL_NAME = "gpt-5.1"

limit_tpm = 30000
try:
    limiter = TPMRateLimiter(tpm_limit=limit_tpm)
except Exception:
    class DummyLimiter:
        def estimate_tokens(self, text): return 0
        def wait_for_tokens(self, tokens): pass
    limiter = DummyLimiter()

def read_context_document(base_dir):
    # Context file: 원주님소개.txt
    context_path = os.path.join(base_dir, "원주님소개.txt")
    try:
        if os.path.exists(context_path):
            with open(context_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            print(f"Warning: Context file not found at {context_path}")
            return ""
    except Exception as e:
        print(f"Error reading context file: {e}")
        return ""

def extract_wonjunim_words_from_json(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        extracted_texts = []
        
        if 'segments' not in data:
            return None

        for segment in data['segments']:
            # Check speaker label. Assuming Speaker 1 is Wonju-nim based on previous script logic
            # "diarization": { "label": "1" }
            speaker_label = segment.get('diarization', {}).get('label')
            
            if speaker_label == '1':
                text = segment.get('text', '').strip()
                if text:
                    extracted_texts.append(text)
        
        if not extracted_texts:
            return None
            
        return " ".join(extracted_texts)
            
    except Exception as e:
        print(f"Error reading/parsing file {file_path}: {e}")
        return None

def refine_text(text, context_content):
    if not client:
        print("OpenAI client is not ready.")
        return None

    prompt = f"""
    [Context Document: 원주님 소개 및 집필 의도]
    {context_content}
    
    [Instruction]
    다음 텍스트는 '원주님'의 말(구어체)을 기록한 것입니다. 위 [Context Document]를 참조하여, 이 말을 책자로 출판하기 위해 다듬어주세요.
    
    지침:
    1. **제목 작성**: 전체 내용을 아우르는 핵심 요약 문장으로 **제목**을 지어주세요. 결과물의 **첫번째 줄**에 `# TITLE: [제목]` 형식으로 작성하세요.
    2. **내용 보존**: 요약하지 말고, 원문의 내용을 **최대한 상세하게** 살리세요. 문맥이나 중요한 뉘앙스, 철학적 사상이 빠지지 않도록 주의하세요.
    3. **문체 변환**: 구어체는 빼고 명확한 **문어체 전문 서적 스타일**로 다듬으세요. (비문, 불필요한 반복, 말더듬 정리, 깊이 있는 어조 유지)
    4. **구조화**: 내용의 흐름에 따라 적절한 **소제목**을 붙여 섹션을 구분하세요. 소제목은 `## [소제목]` 형식을 사용하세요.
    5. **가독성**: 독자가 읽기 편하도록 문단을 적절히 나누세요.

    원문 텍스트:
    {text}
    """

    estimated_tokens = limiter.estimate_tokens(prompt)
    limiter.wait_for_tokens(estimated_tokens)

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are a professional book editor specializing in spiritual and philosophical texts. You refine spoken transcripts into high-quality written text."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return None

def create_docx(content, output_path, original_filename):
    try:
        doc = Document()
        
        # 1. 파일명 기입 (맨 위)
        filename_para = doc.add_paragraph(f"파일명: {original_filename}")
        filename_para.style = 'Normal'
        doc.add_paragraph("") # Spacer

        lines = content.split('\n')
        
        # Check for title
        title_text = "원주님 말씀"
        
        processed_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith("# TITLE:"):
                title_text = line.replace("# TITLE:", "").strip()
                continue
            
            processed_lines.append(line)
        
        # 2. 문서 제목 (요약 문장)
        title_heading = doc.add_heading(title_text, 0)
        title_heading.alignment = 1 # Center
        for run in title_heading.runs:
            run.font.size = Pt(22)

        for line in processed_lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                 doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5

        doc.save(output_path)
        print(f"Successfully saved DOCX to: {output_path}")
        return True
    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return False

def main():
    # Target path setup
    base_dir = r"d:\페르미아\가람컨텐츠팀장"
    source_dir = os.path.join(base_dir, "원주님말씀.json", "2024-2025") # Source JSON files
    target_dir = os.path.join(base_dir, "원주님말씀.doc", "2024-2025.new") # Target DOCX files
    
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        print(f"Created directory: {target_dir}")

    # Load Context
    context_content = read_context_document(base_dir)
    print(f"Loaded Context Document ({len(context_content)} chars)")

    # Find JSON files
    json_files = glob.glob(os.path.join(source_dir, "*.json"))
    
    if not json_files:
        print(f"No JSON files found in {source_dir}")
        return

    print(f"Found {len(json_files)} JSON files to process.")

    # Sort files to process in order
    json_files.sort()

    for json_file in json_files:
        filename = os.path.basename(json_file)
        
        # Filename logic: Base + _문어체.docx
        # Trying to handle common extensions like .mp3.json
        base_name = filename
        if base_name.lower().endswith(".json"):
            base_name = base_name[:-5]
        if base_name.lower().endswith(".mp3"): # If named .mp3.json
            base_name = base_name[:-4]
            
        output_filename = f"{base_name}_문어체.docx"
        output_path = os.path.join(target_dir, output_filename)
        
        # 소제목 반영을 위해 기존 파일이 있어도 건너뛰지 않고 다시 실행하도록 주석 처리
        # if os.path.exists(output_path):
        #     print(f"Skipping {filename} (Output already exists)", flush=True)
        #     continue
        
        print(f"Processing: {filename}", flush=True)
        
        # 1. Extract Wonju-nim's words
        wonjunim_text = extract_wonjunim_words_from_json(json_file)
        if not wonjunim_text:
            print(f"  - No text found for Wonju-nim (Speaker 1) in {filename}", flush=True)
            continue
        
        print(f"  - Extracted {len(wonjunim_text)} chars.", flush=True)

        # 2. Refine
        print("  - Refining (Detailed/Written Style)...", flush=True)
        refined_content = refine_text(wonjunim_text, context_content)
        if not refined_content:
            print("  - Failed to refine.", flush=True)
            continue
            
        # 3. Save
        create_docx(refined_content, output_path, base_name)

if __name__ == "__main__":
    main()
