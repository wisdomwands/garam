import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
output_dir = r"d:\페르미아\가람컨텐츠팀장\원주님말씀.doc\2024-2025.new"
output_filename = "감정기억정보_1_v2.docx"
full_path = os.path.normpath(os.path.join(output_dir, output_filename))
print(f"Saving to: {full_path}")

# Ensure directory exists
os.makedirs(output_dir, exist_ok=True)

# Create Document
doc = Document()

# Title
title = doc.add_heading('감정과 기억, 그리고 에너지의 근원', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Content Data
sections = [
    {
        "header": "1. 기억의 본질과 감정의 힘",
        "content": """뇌의 편도체는 정보를 분류하는 역할을 합니다. 그중에서도 감정을 수반한 기억 정보는 가장 깊이 각인되어 오래도록 지속됩니다. 반면, 감정이 실리지 않은 기억은 쉽게 휘발되어, 의식적으로 유지하려는 노력을 기울이지 않는다면 불과 몇 시간 만에 사라지고 맙니다. 우리가 기억이라고 부르는 것은 정보 그 자체가 아니라, 그 정보에 접근하기 위한 '키워드'일 뿐입니다. 기억 정보 자체는 불안정하기에, 우리는 그 키워드를 재생시키기 위해 끊임없이 정보를 학습하는 것입니다."""
    },
    {
        "header": "2. 감정의 발원지와 언어의 한계",
        "content": """그렇다면 슬픔, 기쁨, 분노와 같은 감정은 어디에서 비롯되는 것일까요? 뇌의 특정 부위를 자극하면 웃거나 눈물을 흘리는 반응이 나타날 수 있습니다. 하지만 뇌는 그 이면에 존재하는 미묘한 감정의 실체와 그 생성 원리까지는 알지 못합니다.
육신을 입고 사는 우리에게 가장 힘겨운 것은 바로 '감정'입니다. 많은 이들이 이 감정 때문에 무너집니다. 그렇다면 이 감정이라는 에너지체는 도대체 어디서 시작되어 위치하는 것일까요?
우리는 흔히 "너 왜 울어? 슬퍼서?"라는 질문에 "응, 슬퍼서 우는 거야"라고 답하며, 자신의 현상을 '슬픔'이라는 언어 정보로 규정해 버립니다. 이것은 일종의 주입식 교육과도 같습니다. 하지만 언어로 정의할 수 없는 미묘한 감정의 영역은 말로 표현하기에 한계가 있습니다."""
    },
    {
        "header": "3. 반응체로서의 육신과 에너지의 작용",
        "content": """인간은 이성이나 정보의 부재 때문이 아니라, 감정의 휘둘림 때문에 "죽겠다" 혹은 "살겠다"를 외칩니다. 그렇다면 감정은 우리 몸의 어느 지점을 장악하고 있는 것일까요? 단순히 신경이 전달되는 통로를 말하는 것이 아닙니다. 자극의 근원이 되어 작용력을 심어주는 그 '감정의 에너지체'가 어디에 존재하는가에 대한 물음입니다.
외부의 소리나 시각 정보는 단지 방아쇠일 뿐입니다. 그것이 내 안에 끓어오르고 있는 특정 에너지체를 건드렸기 때문에 화가 나거나 반응하는 것입니다. 우리는 눈앞의 현상을 보았다고 생각하지만, 실제로는 그 에너지가 끌어당겨진 것입니다."""
    },
    {
        "header": "4. 에너지의 접점과 근원에 대한 탐구",
        "content": """육신은 에너지가 흐르는 통로일 뿐, 그 자체가 원인은 아닙니다. 중요한 것은 스쳐 지나가는 통로가 아니라, 그 에너지를 건드리는 '지점'입니다. 에너지를 감싸고 있는 우리의 육신은 매우 복합적인 반응체입니다. 화를 낸다는 것은 '화'라고 설정된 언어의 껍질이 터져 나오는 것일 뿐, 그 본질적인 작용체(Body)가 어디에 숨어 있는지를 꿰뚫어 보아야 합니다. 이것은 단순히 반응하는 표면이 아닌, 건드려지지 않은 깊은 내면의 영역에 대한 공부가 필요한 이유입니다."""
    }
]

# Write Content
for section in sections:
    # Add Header
    h = doc.add_heading(section["header"], level=1)
    
    # Add Paragraphs
    # Splitting by newline to create separate paragraphs if needed, or just one block
    paragraphs = section["content"].split('\n')
    for p_text in paragraphs:
        if p_text.strip():
            p = doc.add_paragraph(p_text)
            p.paragraph_format.line_spacing = 1.5

# Save
doc.save(full_path)
print(f"File created successfully at: {full_path}")
