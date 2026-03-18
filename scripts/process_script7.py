import os
import glob
import re
from docx import Document
from docx.shared import Pt
from openai import OpenAI
from rate_limiter import TPMRateLimiter

# Initialize OpenAI Client
client = None
try:
    # Key should be set in environment variables
    if os.environ.get("OPENAI_API_KEY"):
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    else:
        print("Warning: OPENAI_API_KEY not found.", flush=True)
except Exception as e:
    print(f"Error initializing OpenAI client: {e}", flush=True)

# Using gpt-5 as requested
MODEL_NAME = "gpt-5" 
limit_tpm = 30000
try:
    limiter = TPMRateLimiter(tpm_limit=limit_tpm)
except Exception:
    # Fallback if TPMRateLimiter is not available or fails
    class DummyLimiter:
        def estimate_tokens(self, text): return 0
        def wait_for_tokens(self, tokens): pass
    limiter = DummyLimiter()

def read_context_document(base_dir):
    context_path = os.path.join(base_dir, "?ì£¼?ìê°?txt")
    try:
        if os.path.exists(context_path):
            with open(context_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            print(f"Warning: Context file not found at {context_path}")
            return ""
    except Exception as e:
        print(f"Error reading context file: {e}")
        return ""

def extract_wonjunim_words_from_txt(file_path):
    """
    Reads a TXT file formatted with 'ë°í??N (Time)' markers.
    Assumes 'ë°í??1' is Wonju-nim.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        extracted_texts = []
        current_speaker = None
        
        # Regex to identify speaker line: e.g., "ë°í??1  (00:03)"
        speaker_pattern = re.compile(r'^ë°í??s+(\d+)\s+')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            match = speaker_pattern.match(line)
            if match:
                current_speaker = match.group(1)
                continue
            
            # If current speaker is 1, collect the text
            # Adjust this logic if Wonju-nim has a different ID in some files,
            # but usually it's 1.
            if current_speaker == '1':
                extracted_texts.append(line)
        
        if not extracted_texts:
            return None
            
        return " ".join(extracted_texts)
            
    except Exception as e:
        print(f"Error reading/parsing file {file_path}: {e}")
        return None

def refine_text(text, context_content):
    if not client:
        print("OpenAI client is not ready.")
        return None

    prompt = f"""
    [Context Document: ?ì£¼???ê° ë°?ì§í ?ë]
    {context_content}
    
    [Instruction]
    ?¤ì ?ì¤?¸ë '?ì£¼????ë§ì? ?´ì©(STT ë°ì·ë³??ë?? ??[Context Document]ë¥?ì°¸ì¡°?ì¬, ??ë§ì???ì±ìë¡?ì¶í?ê¸° ?í´ ?¤ë¬?´ì£¼?¸ì.
    
    ì§ì¹?
    1. **?ëª© ?ì±**: ?ì²´ ?´ì©???ì°ë¥´ë ?µì¬ ?ì½ ë¬¸ì¥?¼ë¡ **?ëª©**??ì§?´ì£¼?¸ì. ê²°ê³¼ë¬¼ì **ì²?ë²ì§¸ ì¤?*??`# TITLE: [?ëª©]` ?ì?¼ë¡ ?ì±?ì¸??
    2. **?´ì© ë³´ì¡´**: ?ì½?ì? ë§ê³ , ?ë¬¸???´ì©??**ìµë????ì¸?ê²** ?´ë¦¬?¸ì. ë¬¸ë§¥?´ë ì¤ì???í?? ì² í???ì?¤ê? ë¹ ì?ì§ ?ëë¡?ì£¼ì?ì¸??
    3. **ë¬¸ì²´ ë³??*: êµ¬ì´ì²´ë? ?ê²© ?ê³  ëªí??**ë¬¸ì´ì²??ë¬¸ ?ì  ?¤í???**ë¡??¤ë¬?¼ì¸?? (ë¹ë¬¸, ë¶í?í ë°ë³µ, ë§ë???ë¦¬, ê¹ì´ ?ë ?´ì¡° ? ì?)
    4. **êµ¬ì¡°??*: **?ì ëª©ì ?°ë¡ ?íì§ ë§ê³ **, ìµë???ë§ì? ê·¸ë?ë¡ë? ë¬¸ë§¥??ë§ê² ë¬¸ì´ì²´ë¡ ë³?í?¸ì. ë¬??ë¥´???ì°?¤ë½ê²??´ì´ì§?ë¡ ?ì ?ì¸?? ë¬¸ë¨? ?´ì©???ë¦??ë°ë????ì ???ë?´ì£¼?¸ì.
    5. **ê°?ì±**: ?ìê° ?½ê¸° ?¸í?ë¡ ë¬¸ë¨???ì ???ë?¸ì.

    ?ë¬¸ ?ì¤??
    {text}
    """

    estimated_tokens = limiter.estimate_tokens(prompt)
    limiter.wait_for_tokens(estimated_tokens)

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are a professional book editor specializing in spiritual and philosophical texts. You refine spoken transcripts into high-quality written text."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return None

def create_docx(content, output_path, original_filename):
    try:
        doc = Document()
        
        # 1. ?ì¼ëª?ê¸°ì (ë§???
        filename_para = doc.add_paragraph(f"?ì¼ëª? {original_filename}")
        filename_para.style = 'Normal'
        doc.add_paragraph("") # Spacer

        lines = content.split('\n')
        
        # Check for title
        title_text = "?ì£¼??ë§ì?"
        
        processed_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith("# TITLE:"):
                title_text = line.replace("# TITLE:", "").strip()
                continue
            
            processed_lines.append(line)
        
        # 2. ë¬¸ì???ëª© (?ì½ ë¬¸ì¥)
        title_heading = doc.add_heading(title_text, 0)
        for run in title_heading.runs:
            run.font.size = Pt(22)

        for line in processed_lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                 doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        print(f"Successfully saved DOCX to: {output_path}")
        return True
    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return False

def main():
    # Target path setup
    base_dir = r"d:\AntiGravity\ÇÑ°¡¶÷ÄÁÅÙÃ÷ÆÀÀå"
    source_dir = os.path.join(base_dir, "2020-2025") # Source TXT files
    target_dir = os.path.join(base_dir, "?ì£¼?ë§?.doc", "2020-2025") # Target DOCX files
    
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        print(f"Created directory: {target_dir}")

    # Load Context
    context_content = read_context_document(base_dir)
    print(f"Loaded Context Document ({len(context_content)} chars)")

    # Find TXT files
    txt_files = glob.glob(os.path.join(source_dir, "*.txt"))
    
    if not txt_files:
        print(f"No TXT files found in {source_dir}")
        return

    print(f"Found {len(txt_files)} TXT files to process.")

    for txt_file in txt_files:
        filename = os.path.basename(txt_file)
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}_ë¬¸ì´ì²?docx"
        output_path = os.path.join(target_dir, output_filename)
        
        if os.path.exists(output_path):
            print(f"Skipping {filename} (Output already exists)", flush=True)
            continue
        
        print(f"Processing: {filename}", flush=True)
        
        # 1. Extract Wonju-nim's words
        wonjunim_text = extract_wonjunim_words_from_txt(txt_file)
        if not wonjunim_text:
            print(f"  - No text found for Wonju-nim (Speaker 1) in {filename}", flush=True)
            continue
        
        print(f"  - Extracted {len(wonjunim_text)} chars.", flush=True)

        # 2. Refine
        print("  - Refining (Detailed/Written Style)...", flush=True)
        refined_content = refine_text(wonjunim_text, context_content)
        if not refined_content:
            print("  - Failed to refine.", flush=True)
            continue
            
        # 3. Save
        create_docx(refined_content, output_path, filename)

if __name__ == "__main__":
    main()
