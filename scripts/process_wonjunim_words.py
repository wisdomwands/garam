import json
import os
import sys
import glob
from docx import Document
from docx.shared import Pt
from openai import OpenAI
from rate_limiter import TPMRateLimiter

# Initialize OpenAI Client
client = None
try:
    if os.environ.get("OPENAI_API_KEY"):
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    else:
        print("Warning: OPENAI_API_KEY not found in environment.")
except Exception as e:
    print(f"Error initializing OpenAI client: {e}")

# Using gpt-5 as requested.
MODEL_NAME = "gpt-5" 
limiter = TPMRateLimiter(tpm_limit=30000)

def read_context_document(base_dir):
    context_path = os.path.join(base_dir, "?ì£¼?ìê°?txt")
    try:
        if os.path.exists(context_path):
            with open(context_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            print(f"Warning: Context file not found at {context_path}")
            return ""
    except Exception as e:
        print(f"Error reading context file: {e}")
        return ""

def extract_wonjunim_words(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
            
            extracted_texts = []
            
            if 'segments' in data:
                for segment in data['segments']:
                    # Check for Speaker 1 or Name A (Wonju-nim)
                    speaker = segment.get('speaker', {})
                    label = segment.get('diarization', {}).get('label')
                    
                    is_wonjunim = False
                    if speaker.get('name') == 'A':
                        is_wonjunim = True
                    elif speaker.get('label') == '1':
                        is_wonjunim = True
                    elif label == '1':
                        is_wonjunim = True
                        
                    if is_wonjunim:
                        extracted_texts.append(segment['text'])
            
            if not extracted_texts:
                return None
                
            return " ".join(extracted_texts)
            
    except Exception as e:
        print(f"Error reading/parsing file {file_path}: {e}")
        return None

def refine_text(text, context_content):
    if not client:
        print("OpenAI client is not ready.")
        return None

    prompt = f"""
    [Context Document: ?ì£¼???ê° ë°?ì§í ?ë]
    {context_content}
    
    [Instruction]
    ?¤ì ?ì¤?¸ë '?ì£¼????ë§ì? ?´ì©(STT ë°ì·ë³??ë?? ??[Context Document]ë¥?ì°¸ì¡°?ì¬, ??ë§ì???ì±ìë¡?ì¶í?ê¸° ?í´ ?¤ë¬?´ì£¼?¸ì.
    
    ì§ì¹?
    1. **?ëª© ?ì±**: ?ì²´ ?´ì©???ì°ë¥´ë ?µì¬ ?ì½ ë¬¸ì¥?¼ë¡ **?ëª©**??ì§?´ì£¼?¸ì. ê²°ê³¼ë¬¼ì **ì²?ë²ì§¸ ì¤?*??`# TITLE: [?ëª©]` ?ì?¼ë¡ ?ì±?ì¸??
    2. **?´ì© ë³´ì¡´**: ?ì½?ì? ë§ê³ , ?ë¬¸???´ì©??**ìµë????ì¸?ê²** ?´ë¦¬?¸ì. ë¬¸ë§¥?´ë ì¤ì???í?? ì² í???ì?¤ê? ë¹ ì?ì§ ?ëë¡?ì£¼ì?ì¸??
    3. **ë¬¸ì²´ ë³??*: êµ¬ì´ì²´ë? ?ê²© ?ê³  ëªí??**ë¬¸ì´ì²??ë¬¸ ?ì  ?¤í???**ë¡??¤ë¬?¼ì¸?? (ë¹ë¬¸, ë¶í?í ë°ë³µ, ë§ë???ë¦¬, ê¹ì´ ?ë ?´ì¡° ? ì?)
    4. **êµ¬ì¡°??*: ?´ì©???¼ë¦¬?ì¸ ?ë¦???°ë¼ **[?ì£¼ì ]**, **[?ì£¼??**ë¥?ë¶ì¬??ê³ì¸µ?ì¼ë¡????ë¦¬?ì¸??
    5. **ê°?ì±**: ?ìê° ?½ê¸° ?¸í?ë¡ ë¬¸ë¨???ì ???ë?¸ì.

    ?ë¬¸ ?ì¤??
    {text}
    """

    # Truncate text if too long to fit in context with output
    # Using simple length check for safety, assuming 1 token ~ 4 chars (Korean ~1-2 chars). 
    # GPT-4o context is large, GPT-5 presumably too. 
    
    estimated_tokens = limiter.estimate_tokens(prompt)
    limiter.wait_for_tokens(estimated_tokens)

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are a professional book editor specializing in spiritual and philosophical texts. You refine spoken transcripts into high-quality written text."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return None

def create_docx(content, output_path, original_filename):
    try:
        doc = Document()
        
        # 1. ?ì¼ëª?ê¸°ì (ë§???
        filename_para = doc.add_paragraph(f"?ì¼ëª? {original_filename}")
        filename_para.style = 'Normal' # Or a smaller style if preferred
        doc.add_paragraph("") # Spacer

        lines = content.split('\n')
        
        # Check for title
        title_text = "?ì£¼??ë§ì?"
        
        processed_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith("# TITLE:"):
                title_text = line.replace("# TITLE:", "").strip()
                continue # Do not add title line to body yet, add as Heading 0
            
            processed_lines.append(line)

        # 2. ë¬¸ì???ëª© (?ì½ ë¬¸ì¥)
        title_heading = doc.add_heading(title_text, 0)
        for run in title_heading.runs:
            run.font.size = Pt(22)

        for line in processed_lines:
            if line.startswith('# '): # Remaining H1s become H1 or H2 depending on structure
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                 doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        print(f"Successfully saved DOCX to: {output_path}")
        return True
    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return False

def main():
    # Target path: .\?ì£¼?ë§?.json\2013
    base_dir = r"d:\AntiGravity\ÇÑ°¡¶÷ÄÁÅÙÃ÷ÆÀÀå"
    source_dir = os.path.join(base_dir, "?ì£¼?ë§?.json", "2013ë¹ë??)
    target_dir = os.path.join(base_dir, "?ì£¼?ë§?.doc", "2013ë¹ë??)
    
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        print(f"Created directory: {target_dir}")

    # Load Context
    context_content = read_context_document(base_dir)
    print(f"Loaded Context Document ({len(context_content)} chars)")

    # Find JSON files
    json_files = glob.glob(os.path.join(source_dir, "*.json"))
    
    if not json_files:
        print(f"No JSON files found in {source_dir}")
        return

    print(f"Found {len(json_files)} JSON files to process.")

    for json_file in json_files:
        filename = os.path.basename(json_file)
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}_ë¬¸ì´ì²?docx"
        output_path = os.path.join(target_dir, output_filename)
        
        if os.path.exists(output_path):
            print(f"Skipping {filename} (Output already exists)")
            continue
        
        print(f"Processing: {filename}")
        
        # 1. Extract Wonju-nim's words
        wonjunim_text = extract_wonjunim_words(json_file)
        if not wonjunim_text:
            print(f"  - No text found for Wonju-nim in {filename}")
            continue
            
        # 2. Refine
        print("  - Refining (Detailed/Written Style)...")
        refined_content = refine_text(wonjunim_text, context_content)
        if not refined_content:
            print("  - Failed to refine.")
            continue
            
        # 3. Save
        create_docx(refined_content, output_path, filename)

if __name__ == "__main__":
    main()
