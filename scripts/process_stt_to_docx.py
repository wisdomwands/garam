import json
import os
import sys
from docx import Document
from openai import OpenAI
from rate_limiter import TPMRateLimiter

# Initialize OpenAI Client
client = None
try:
    # Using environment variable
    if os.environ.get("OPENAI_API_KEY"):
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    else:
        print("Warning: OPENAI_API_KEY not found.")
except Exception as e:
    print(f"Error initializing OpenAI client: {e}")

MODEL_NAME = "gpt-4o"
limiter = TPMRateLimiter(tpm_limit=30000)

def extract_text_from_json(file_path):
    try:
        # Try reading as JSON first
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            try:
                data = json.load(f)
                # Check standard Clova format
                if 'text' in data:
                    return data['text']
                # If segments exist but no top-level text (rare for Clova fullText:True but possible)
                elif 'segments' in data:
                     return " ".join([seg['text'] for seg in data['segments']])
                else:
                    return None
            except json.JSONDecodeError:
                # If JSON parsing fails, try reading as plain text
                print(f"File {file_path} is not valid JSON. Attempting to read as plain text.")
                f.seek(0)
                return f.read()
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return None
    except Exception as e:
        print(f"Error reading JSON file {file_path}: {e}")
        return None

def refine_text_style(text):
    if not client:
        print("OpenAI client is not ready.")
        return None

    prompt = """
    ?¤ì ?ì¤?¸ë ?ì± ?¸ì(STT) ê²°ê³¼ë¬¼ì?ë¤. ???´ì©??ë°í?¼ë¡ '?ì± ë¶ì¼???ë¬¸ êµì¬' ?íë¡?ì±ìë¥?ë§ë¤?¤ê³  ?©ë?? 
    ?¤ì ì§ì¹¨ì ?°ë¼ ?ì¤?¸ë? ë¬¸ì´ì²´ë¡ ë³?íê³??¬êµ¬?±í´ì£¼ì¸??

    1. **ë¬¸ì²´**: ê¹ì´ ?ê³  ê¶ì ?ì¼ë©´ì???´í´?ê¸° ?¬ì´ ë¬¸ì´ì²??ì  ?¤í???ë¡?ë³?í?¸ì. êµ¬ì´ì²´ì êµ°ë?ê¸°, ë°ë³µ, ë¹ë¬¸? ?ê±°?ê±°???ì ?ì¸??
    2. **êµ¬ì¡°??*: ?´ì©???¼ë¦¬?ì¸ ?ë¦???°ë¼ ?¬ë°°ì¹íê³? ?ì ??[?ì£¼ì ], [?ì£¼??ë¥?ë¶ì¬??ê³ì¸µ?ì¼ë¡?êµ¬ì¡°?í´ì£¼ì¸??
    3. **?´ì© ë³´ì**: ë¬¸ë§¥??ë¹ ì§ ?ì?¬ë ì£¼ì´ë¥?ë³´ì?ì¬ ë¬¸ì¥??ë§¤ë?½ê² ?´ì´ì§?ë¡ ?ì¸?? ?ë?ê° ëªí?ì? ?ì? ë¶ë¶ì? ë¬¸ë§¥???µí´ ?©ë¦¬?ì¼ë¡?ì¶ë¡ ?ì¬ ?¤ë¬?´ì£¼?¸ì.
    4. **?ì**: ìµì¢ ê²°ê³¼ë¬¼ì? ê·¸ë?ë¡?ë³µì¬?´ì ì±ì ?£ì ???ë ?ì±???ê³  ?í?¬ì¼ ?©ë?? 

    ?ë¬¸ ?ì¤??
    """ + text

    estimated_tokens = limiter.estimate_tokens(prompt)
    limiter.wait_for_tokens(estimated_tokens)

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are a professional editor specializing in spiritual and philosophical texts."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return None

def create_docx(content, output_path):
    try:
        doc = Document()
        doc.add_heading('?ì± êµì¬ ?ê³ ', 0)

        # Split content by lines to process basic formatting if needed
        # For now, just adding paragraphs. GPT output typically has markdown headers like ## or ###.
        # We can do a simple parsing to convert Markdown headers to Docx headers.
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                 doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        print(f"Successfully saved DOCX to: {output_path}")
        return True
    except Exception as e:
        print(f"Error creating DOCX file: {e}")
        return False

def main():
    target_file = r"d:\AntiGravity\ÇÑ°¡¶÷ÄÁÅÙÃ÷ÆÀÀå\sample.txt"
    if len(sys.argv) > 1:
        target_file = sys.argv[1]

    print(f"Processing: {target_file}")
    
    # 1. Extract
    original_text = extract_text_from_json(target_file)
    if not original_text:
        print("Failed to extract text from JSON.")
        return

    print(f"Extracted {len(original_text)} characters.")

    # 2. Refine
    print("Refining text with OpenAI...")
    refined_content = refine_text_style(original_text)
    if not refined_content:
        print("Failed to refine text.")
        return

    # 3. Save to Docx
    base_name = os.path.splitext(target_file)[0]
    output_docx = f"{base_name}_booklet.docx"
    
    create_docx(refined_content, output_docx)

if __name__ == "__main__":
    main()
