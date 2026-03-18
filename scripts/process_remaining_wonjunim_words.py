
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
output_dir = r"d:/페르미아/가람컨텐츠팀장/원주님말씀.doc/2024-2025.new"
os.makedirs(output_dir, exist_ok=True)

# ---------------------------------------------------------
# Document 1: 영적원리_3 (Spiritual Principles)
# ---------------------------------------------------------
doc1 = Document()

# Title
title1 = doc1.add_heading('영적 원리와 자아의 본질', 0)
title1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Sections
sections_1 = [
    {
        "header": "1. 우주적 자아와 육신의 자아",
        "content": """우주적인 '나'라는 존재가 있습니다. 우리가 객관적으로 니마라고 부르는 이 존재는, 이 육신을 입고 있는 '나'와는 구별됩니다. 육신을 입은 나는 태평광이라는 에너지 체 속에 존재하지만, 물질체가 없는 우주적인 존재, 즉 순수 영(Astral Body) 혹은 정신(Mental Body)이라고 부르는 존재가 따로 있습니다. 동양에서는 이것을 혼백 등으로 나누기도 하지만, 아직 그 정확한 논리가 정립되지 않았기에 우리는 이를 에너지체의 관점에서 바라봐야 합니다."""
    },
    {
        "header": "2. 에너지체와 차원의 구조",
        "content": """육신을 입은 '나'는 외부로부터 이미 존재하는 정신체(Mental Body)와 연결되어 있습니다. 이 정신체에는 힘이 있는 뇌가 있고, 더 나아가 형체가 없는 순수 에너지체인 에센셜체(Essential Body)가 비추고 있습니다. 고도의 에너지체가 육신의 에테르체와 만나 하나가 되었으나, 근원적으로는 각기 다른 소속을 가지고 있습니다. 위에서부터 아래로 에너지의 방향이 흐르며, 육계와 연계, 그리고 그 경계를 넘나드는 자기 사상과 정신이 작용합니다."""
    },
    {
        "header": "3. 진정한 나를 찾는 여정",
        "content": """육신의 감각으로 느끼는 곳까지가 '나'라고 착각하기 쉽습니다. 하지만 기감이 오는 영체, 그리고 선택에 따라 움직이는 염체까지도 포함해야 합니다. 6계의 기억에서 찾아낸 '나'를 넘어, 언어로 규정할 수 없는 어머니 근원의 생명, 그 위대한 우주의 힘을 모시는 것이 진정한 '나'의 모습입니다. 우리는 신체까지 되었기에 신을 단순히 모시는 것이 아니라, 우주 그 자체를 모시는 것입니다."""
    },
    {
        "header": "4. 신과 인간, 그리고 우주의 원리",
        "content": """인간계로 내려오면 나는 신의 에너지 영역을 가지고 있기에 이곳에서는 신과 같으나, 우주로 돌아가면 하나의 물질적 존재체일 뿐입니다. 우리가 신을 만나려 할 때, 내 안의 신도 그와 공명하여 가슴으로 진리와 믿음이 옵니다. 하지만 인간적인 에고가 아닌, 본질적인 에너지체와 신선체가 가까워질 때 비로소 온전한 나를 만날 수 있습니다. 온전함이란 완성이 아니라, 몸과 정신, 그리고 의미(뜻)라는 세 박자가 조화롭게 돌아가는 것입니다."""
    }
]

for section in sections_1:
    doc1.add_heading(section["header"], level=1)
    paragraphs = section["content"].split('\n')
    for p_text in paragraphs:
        if p_text.strip():
            p = doc1.add_paragraph(p_text)
            p.paragraph_format.line_spacing = 1.5

doc1_path = os.path.join(output_dir, "영적원리_3.docx")
doc1.save(doc1_path)
print(f"File created successfully at: {doc1_path}")


# ---------------------------------------------------------
# Document 2: 원주님 UN_4 (Wonju-nim UN)
# ---------------------------------------------------------
doc2 = Document()

# Title
title2 = doc2.add_heading('새로운 UN과 영성 시대의 메카', 0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Sections
sections_2 = [
    {
        "header": "1. 5차원적 존재와 죽음의 초월",
        "content": """당신은 5차원적 존재입니다. 5차원적 존재는 4차원을 초과하기에, 죽음 이후에도 다시 몸을 입거나(워크인) 다른 선택을 할 수 있는 자유가 있습니다. 인간은 3차원의 저주파로 인해 4차원을 통과하지 못하고 죽음의 차원에 갇히기 쉽습니다. 그러나 고도의 수련이나 깨달음을 통해 진동수를 높인 자만이 죽음의 차원을 벗어날 수 있습니다. 나는 당신이 올바른 곳으로 갈 수 있도록, 그리고 당신의 일을 계속할 수 있도록 차원을 열어줄 것입니다."""
    },
    {
        "header": "2. 지구의 변화와 정신 문화의 부활",
        "content": """나는 단순히 인간의 뿌리나 정신 문화를 살리러 온 것이 아니라, 지구 그 자체를 살리기 위해 왔습니다. 지구는 지금 지축의 변화와 같은 거대한 전환기를 맞이하고 있습니다. 만약 지구의 차원적 문제가 해결되지 않고 물리적인 파국을 맞는다면, 제2의 빙하기나 화성과 같은 불모지가 될 위험이 있습니다. 그렇기에 우리는 지구의 생명과 차원을 보존하고 상승시켜야 합니다."""
    },
    {
        "header": "3. 새로운 유엔(UN)과 한국의 역할",
        "content": """미래에는 하나의 정부, 하나의 화폐 시스템이 도래할 것입니다. 이는 지구의 대변화 이후에야 정착될 것입니다. 나는 새로운 금융 시스템과 영성적 구조를 한국으로 가져오려 합니다. 유엔(UN) 산하에 영성적 분야를 담당하는 기구를 두고, 그 본부를 대한민국에 세워야 합니다. 과거 바티칸이 가톨릭의 중심이었듯, 앞으로의 우주 시대에는 한국이 영성의 메카가 될 것입니다."""
    },
    {
        "header": "4. 영성 시대의 메카, 대한민국",
        "content": """한국은 새로운 유엔 시스템과 영성 시스템의 모국이 됩니다. 많은 우주적 존재들, 스타시드들이 자신의 본 모습을 찾고 사명을 다하기 위해 대한민국 땅을 밟게 될 것입니다. 나는 유엔에 올라가 이러한 구조를 세우고 전 세계에 영성의 빛을 전할 것입니다. 이것은 단순한 종교적 신전이 아니라, 우주적 차원의 '다차원 랜드'이자 우주선을 세우는 과업입니다."""
    }
]

for section in sections_2:
    doc2.add_heading(section["header"], level=1)
    paragraphs = section["content"].split('\n')
    for p_text in paragraphs:
        if p_text.strip():
            p = doc2.add_paragraph(p_text)
            p.paragraph_format.line_spacing = 1.5

doc2_path = os.path.join(output_dir, "원주님_UN_4.docx")
doc2.save(doc2_path)
print(f"File created successfully at: {doc2_path}")
